#!/usr/bin/python
#-*- coding: latin-1 -*-

print("###############################################")
print("############ LEYENDO DOCUMENTO WORD ###########")
print("###############################################")

from docx import Document
##################################################################
texto = []
i = 0
document = Document("C:/Users/admin/Projects/Documentos/Abstrct.docx")
#document.paragraphs
for texto in document.paragraphs:
    if (len(texto.text)>0):
        print("\n")
        print(texto.text)


#input("\nPrecione una tecla para continuar...")
